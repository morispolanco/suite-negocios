import streamlit as st
import requests
import json
from io import BytesIO
from docx import Document

# Configuración de la página principal
st.set_page_config(page_title="Suite de Aplicaciones de Negocios", page_icon="💼", layout="wide")

# Selección de la aplicación en la suite
app_selection = st.sidebar.selectbox("Selecciona una aplicación", [
    "Generación de Plan de Negocios", 
    "Generación de Propuestas de Negocios",
    "Análisis de Mercado",
    "Desarrollo de Estrategia Empresarial",
    "Proyecciones Financieras"
])

# Función para buscar datos de mercado usando la API de Serper
def buscar_informacion(query, country):
    url = "https://google.serper.dev/search"
    payload = json.dumps({
        "q": f"{query} market data, trends, competitors in {country}",
        "num": 10
    })
    headers = {
        'X-API-KEY': st.secrets["SERPER_API_KEY"],
        'Content-Type': 'application/json'
    }
    response = requests.post(url, headers=headers, data=payload)
    return response.json()

# Función para generar contenido usando la API de Together
def generar_contenido(prompt):
    url = "https://api.together.xyz/inference"
    payload = json.dumps({
        "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
        "prompt": prompt,
        "max_tokens": 2048,
        "temperature": 0.2,
        "top_p": 0.9,
        "top_k": 50,
        "repetition_penalty": 1.1,
        "stop": ["Resumen:", "Puntos clave:"]
    })
    headers = {
        'Authorization': f'Bearer {st.secrets["TOGETHER_API_KEY"]}',
        'Content-Type': 'application/json'
    }
    response = requests.post(url, headers=headers, data=payload)
    return response.json()['output']['choices'][0]['text'].strip()

# Función para crear un archivo DOCX
def create_docx(titulo, secciones):
    doc = Document()
    doc.add_heading(titulo, 0)
    for heading, content in secciones.items():
        doc.add_heading(heading, level=1)
        doc.add_paragraph(content)
    return doc

# Aplicación: Generación de Plan de Negocios
if app_selection == "Generación de Plan de Negocios":
    st.title("Generación de Plan de Negocios")
    idea_negocio = st.text_input("Ingresa tu idea de negocio o sector:")
    pais = st.text_input("Ingresa el país:")

    if st.button("Generar Plan de Negocios"):
        if idea_negocio and pais:
            with st.spinner("Buscando datos de mercado y generando plan..."):
                resultados_busqueda = buscar_informacion(idea_negocio, pais)

                st.write("Datos crudos obtenidos de la búsqueda del Serper API:")
                st.json(resultados_busqueda)  # Debug step

                plan_negocio = ""

                for item in resultados_busqueda.get("organic", []):
                    snippet = item.get("snippet", "")
                    plan_negocio += generar_contenido(f"Genera un plan de negocios basado en la siguiente información en {pais}: {snippet}")

                # Mostrar el plan de negocios
                st.subheader("Plan de Negocios Generado")
                st.markdown(plan_negocio)

                # Botón para descargar el plan de negocios
                doc = create_docx("Plan de Negocios", {"Plan de Negocios": plan_negocio})
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Descargar Plan de Negocios en DOCX",
                    data=buffer,
                    file_name=f"Plan_Negocios_{idea_negocio.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Por favor, ingresa una idea de negocio y un país.")

# Aplicación: Generación de Propuestas de Negocios
elif app_selection == "Generación de Propuestas de Negocios":
    st.title("Generación de Propuestas de Negocios")
    idea_propuesta = st.text_input("Ingresa tu idea de propuesta o sector:")
    pais = st.text_input("Ingresa el país:")

    if st.button("Generar Propuesta de Negocios"):
        if idea_propuesta and pais:
            with st.spinner("Buscando datos de mercado y generando propuesta..."):
                resultados_busqueda = buscar_informacion(idea_propuesta, pais)

                st.write("Datos crudos obtenidos de la búsqueda del Serper API:")
                st.json(resultados_busqueda)  # Debug step

                propuesta_negocio = ""

                for item in resultados_busqueda.get("organic", []):
                    snippet = item.get("snippet", "")
                    propuesta_negocio += generar_contenido(f"Genera una propuesta de negocio basada en la siguiente información en {pais}: {snippet}")

                # Mostrar la propuesta de negocio
                st.subheader("Propuesta de Negocios Generada")
                st.markdown(propuesta_negocio)

                # Botón para descargar la propuesta de negocios
                doc = create_docx("Propuesta de Negocios", {"Propuesta de Negocios": propuesta_negocio})
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Descargar Propuesta de Negocios en DOCX",
                    data=buffer,
                    file_name=f"Propuesta_Negocios_{idea_propuesta.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Por favor, ingresa una idea de propuesta y un país.")

# Aplicación: Análisis de Mercado
elif app_selection == "Análisis de Mercado":
    st.title("Análisis de Mercado")
    mercado = st.text_input("Ingresa el mercado o sector a analizar:")
    pais = st.text_input("Ingresa el país:")

    if st.button("Realizar Análisis de Mercado"):
        if mercado and pais:
            with st.spinner("Buscando datos de mercado y generando análisis..."):
                resultados_busqueda = buscar_informacion(mercado, pais)

                st.write("Datos crudos obtenidos de la búsqueda del Serper API:")
                st.json(resultados_busqueda)  # Debug step

                analisis_mercado = ""

                for item in resultados_busqueda.get("organic", []):
                    snippet = item.get("snippet", "")
                    analisis_mercado += generar_contenido(f"Realiza un análisis de mercado basado en la siguiente información en {pais}: {snippet}")

                # Mostrar el análisis de mercado
                st.subheader("Análisis de Mercado Generado")
                st.markdown(analisis_mercado)

                # Botón para descargar el análisis de mercado
                doc = create_docx("Análisis de Mercado", {"Análisis de Mercado": analisis_mercado})
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Descargar Análisis de Mercado en DOCX",
                    data=buffer,
                    file_name=f"Analisis_Mercado_{mercado.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Por favor, ingresa un mercado y un país.")

# Aplicación: Desarrollo de Estrategia Empresarial
elif app_selection == "Desarrollo de Estrategia Empresarial":
    st.title("Desarrollo de Estrategia Empresarial")
    estrategia = st.text_input("Ingresa el área de estrategia a desarrollar:")
    pais = st.text_input("Ingresa el país:")

    if st.button("Desarrollar Estrategia Empresarial"):
        if estrategia and pais:
            with st.spinner("Buscando información y desarrollando estrategia..."):
                resultados_busqueda = buscar_informacion(estrategia, pais)

                st.write("Datos crudos obtenidos de la búsqueda del Serper API:")
                st.json(resultados_busqueda)  # Debug step

                estrategia_empresarial = ""

                for item in resultados_busqueda.get("organic", []):
                    snippet = item.get("snippet", "")
                    estrategia_empresarial += generar_contenido(f"Desarrolla una estrategia empresarial basada en la siguiente información en {pais}: {snippet}")

                # Mostrar la estrategia empresarial
                st.subheader("Estrategia Empresarial Generada")
                st.markdown(estrategia_empresarial)

                # Botón para descargar la estrategia empresarial
                doc = create_docx("Estrategia Empresarial", {"Estrategia Empresarial": estrategia_empresarial})
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Descargar Estrategia Empresarial en DOCX",
                    data=buffer,
                    file_name=f"Estrategia_Empresarial_{estrategia.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Por favor, ingresa un área de estrategia y un país.")

# Aplicación: Proyecciones Financieras
elif app_selection == "Proyecciones Financieras":
    st.title("Proyecciones Financieras")
    empresa = st.text_input("Ingresa el nombre de la empresa o sector:")
    pais = st.text_input("Ingresa el país:")

    if st.button("Generar Proyecciones Financieras"):
        if empresa and pais:
            with st.spinner("Buscando información financiera y generando proyecciones..."):
                resultados_busqueda = buscar_informacion(empresa, pais)

                st.write("Datos crudos obtenidos de la búsqueda del Serper API:")
                st.json(resultados_busqueda)  # Debug step

                proyecciones_financieras = ""

                for item in resultados_busqueda.get("organic", []):
                    snippet = item.get("snippet", "")
                    proyecciones_financieras += generar_contenido(f"Genera proyecciones financieras para la empresa o sector basado en la siguiente información en {pais}: {snippet}")

                # Mostrar las proyecciones financieras
                st.subheader("Proyecciones Financieras Generadas")
                st.markdown(proyecciones_financieras)

                # Botón para descargar las proyecciones financieras
                doc = create_docx("Proyecciones Financieras", {"Proyecciones Financieras": proyecciones_financieras})
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Descargar Proyecciones Financieras en DOCX",
                    data=buffer,
                    file_name=f"Proyecciones_Financieras_{empresa.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Por favor, ingresa el nombre de la empresa o sector y un país.")
